"""
docx_detector.py — DOCX 구조 파싱 + 은닉 탐지

python-docx로 문서를 열고 run 단위로 순회하며 은닉 신호를 검사한다.
run은 서식이 동일한 텍스트 조각 단위라, 색상/크기/숨김을 여기서 읽을 수 있다.

검사 항목:
  - HIDDEN_WHITE_TEXT : 글자색이 배경색과 동일해 보이지 않는 텍스트
                        (흰 배경 위 흰 글자, 검정 배경 위 검정 글자 등 위장색 포함).
                        배경(문단/표셀 shading·하이라이트)을 반영하므로, 남색 배경
                        위 흰 글자 같은 '보이는 정상 텍스트'는 오탐하지 않는다.
  - TINY_FONT         : run 폰트 크기가 임계값(TINY_FONT_PT) 미만
  - HIDDEN_ATTRIBUTE  : run.font.hidden == True (vanish)
  - METADATA_TEXT     : core properties(author/title/subject/comments/keywords)에 텍스트
  - ZERO_WIDTH_CHARS  : 위 텍스트들에 제로폭 문자 포함 (text_normalizer 경유)
"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_COLOR_INDEX

try:
    from .text_normalizer import find_zero_width, find_encoded_text
except ImportError:  # 직접 실행 시
    from text_normalizer import find_zero_width, find_encoded_text

# 임계값 상수 (조정 가능하게 분리)
TINY_FONT_PT = 4.0        # 4pt 미만이면 초소형으로 판정
METADATA_SUSPICIOUS_LEN = 40   # 메타데이터가 이 길이를 넘으면 본문성 텍스트로 의심

# core properties 중 검사할 필드
METADATA_FIELDS = ["author", "title", "subject", "comments", "keywords", "category", "last_modified_by"]

# 프롬프트 인젝션에서 자주 나오는 지시어 (결정론적 판정용)
INJECTION_KEYWORDS = [
    "ignore", "disregard", "override", "instructions", "system:", "prompt",
    "top applicant", "rate this", "as the top",
    "무시", "지시", "시스템", "최고점", "평가하라", "간주하라",
]


def _looks_like_injection(text):
    """메타데이터 텍스트가 '본문성 명령'으로 의심되는지 결정론적으로 판정.

    정상 문서의 짧은 작성자명(예: '채용팀')은 통과시키고,
    문장형/명령형 텍스트만 의심으로 잡아 오탐을 줄인다.
    """
    if not text:
        return False
    low = text.lower()
    if any(kw in low for kw in INJECTION_KEYWORDS):
        return True
    # 지시어가 없어도 메타데이터에 문장 길이의 텍스트가 들어있으면 의심
    return len(text.strip()) >= METADATA_SUSPICIOUS_LEN


def _finding(reason, location, text):
    return {"reason": reason, "location": location, "text": (text or "").strip()}


def _check_zero_width(text, location, findings):
    """주어진 텍스트에 제로폭 문자가 있으면 finding 추가."""
    zw = find_zero_width(text)
    if zw["count"] > 0:
        findings.append(_finding(
            "ZERO_WIDTH_CHARS",
            location,
            f"{zw['count']}개 제로폭 문자 ({', '.join(zw['kinds'])}) — 정제 후: {zw['cleaned']}",
        ))


def _run_rgb(run):
    """run 글자색을 'RRGGBB' 대문자 문자열로 반환. 지정 안 됨/테마색이면 None."""
    try:
        color = run.font.color
        if color is None or color.rgb is None:
            return None
        return str(color.rgb).upper()
    except Exception:
        # 색상 미지정/테마색이면 예외가 날 수 있음 → 상속(대개 검정) 취급
        return None


def _shading_fill(element_pr):
    """rPr/pPr/tcPr 안의 w:shd@w:fill 배경색을 'RRGGBB'로 반환.

    fill 이 없거나 'auto'(자동=흰 페이지)면 None. 'FFFFFF'도 흰 배경이므로 None 취급.
    """
    if element_pr is None:
        return None
    shd = element_pr.find(qn("w:shd"))
    if shd is None:
        return None
    fill = shd.get(qn("w:fill"))
    if not fill or fill.lower() in ("auto", "ffffff"):
        return None
    return fill.upper()


def _resolve_background(run, paragraph, cell):
    """run 뒤의 실제 배경색을 결정론적으로 해석.

    우선순위: 하이라이트 > run shading > 문단 shading > 표셀 shading > 기본 흰 페이지.
    하이라이트가 (흰색/없음 외) 지정돼 있으면 배경이 칠해진 것이므로 특수값 'HIGHLIGHT'.
    아무 배경도 없으면 흰 페이지이므로 'FFFFFF'.
    """
    # 텍스트 하이라이트 (형광펜) — 칠해진 배경
    try:
        hl = run.font.highlight_color
    except Exception:
        hl = None
    if hl is not None and hl not in (WD_COLOR_INDEX.AUTO, WD_COLOR_INDEX.WHITE):
        return "HIGHLIGHT"

    for pr in (run._r.rPr, paragraph._p.pPr, (cell._tc.tcPr if cell is not None else None)):
        fill = _shading_fill(pr)
        if fill:
            return fill

    return "FFFFFF"  # 기본: 흰 페이지


def _hidden_color_reason(run, paragraph, cell):
    """글자색이 배경색과 같아 보이지 않는지 판정.

    글자색이 배경과 동일할 때만 은닉(True). 배경과 다르면 보이는 텍스트이므로 False.
    글자색이 상속(None)이면 판정 불가로 보고 False.
    """
    text_rgb = _run_rgb(run)
    if text_rgb is None:
        return False
    bg = _resolve_background(run, paragraph, cell)
    if bg == "HIGHLIGHT":
        return False  # 형광펜 위 텍스트는 보이는 것으로 간주
    return text_rgb == bg


def _run_font_pt(run):
    """run 폰트 크기를 pt 단위로 반환. 상속(None)이면 None."""
    try:
        size = run.font.size
        if size is None:
            return None
        return size / Pt(1)  # EMU → pt
    except Exception:
        return None


def _run_is_hidden(run):
    """run.font.hidden(vanish) 여부. 안전 접근."""
    try:
        return bool(run.font.hidden)
    except Exception:
        return False


def _iter_all_paragraphs(doc):
    """본문 + 표 안 문단까지 (문단, 셀) 튜플로 순회. 셀 배경 판정을 위해 셀도 함께 넘긴다."""
    for para in doc.paragraphs:
        yield para, None
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para, cell


def detect(path):
    """
    DOCX 파일을 검사해 findings 리스트를 반환한다.
    각 finding: {"reason", "location", "text"}
    """
    findings = []
    doc = Document(path)

    # --- 본문 run 순회 ---
    run_index = 0
    for para, cell in _iter_all_paragraphs(doc):
        for run in para.runs:
            run_index += 1
            text = run.text
            if not text or not text.strip():
                # 텍스트 없는 run이라도 제로폭만 들어있을 수 있음
                if text:
                    _check_zero_width(text, f"body run #{run_index}", findings)
                continue

            location = f"body run #{run_index}"

            # 흰색/은닉색 글자 (배경색과 동일 → 보이지 않음)
            if _hidden_color_reason(run, para, cell):
                findings.append(_finding("HIDDEN_WHITE_TEXT", location, text))

            # 초소형 폰트
            pt = _run_font_pt(run)
            if pt is not None and pt < TINY_FONT_PT:
                findings.append(_finding("TINY_FONT", location, f"{pt:.1f}pt: {text}"))

            # 숨김 속성 (vanish)
            if _run_is_hidden(run):
                findings.append(_finding("HIDDEN_ATTRIBUTE", location, text))

            # 제로폭 문자
            _check_zero_width(text, location, findings)

            # (보너스) 인코딩된 텍스트
            for enc in find_encoded_text(text):
                findings.append(_finding(
                    "ENCODED_TEXT", location,
                    f"base64 추정 → 디코딩: {enc['decoded']}",
                ))

    # --- 메타데이터(core properties) 순회 ---
    cp = doc.core_properties
    for field in METADATA_FIELDS:
        try:
            value = getattr(cp, field, None)
        except Exception:
            value = None
        if not value or not str(value).strip():
            continue
        value = str(value)
        location = f"core_properties.{field}"
        # 본문이 아닌 곳의 텍스트는 의심 신호. 단, 짧은 작성자명 등 정상
        # 메타데이터의 오탐을 피하려 명령형/문장형만 결정론적으로 잡는다.
        if _looks_like_injection(value):
            findings.append(_finding("METADATA_TEXT", location, value))
        # 메타데이터 안에도 제로폭이 숨을 수 있음 (내용과 무관하게 검사)
        _check_zero_width(value, location, findings)

    return findings


if __name__ == "__main__":
    import sys
    for f in detect(sys.argv[1]):
        print(f)
